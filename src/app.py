import sys
import os
from pathlib import Path

sys.path.append(str(Path(__file__).parent.parent))

import streamlit as st
import pdfplumber
import io
import re
import time
import logfire
from docx import Document
from dotenv import load_dotenv

load_dotenv(override=True)
logfire.configure(token=os.getenv("LOGFIRE_TOKEN"))

from src.agent.agent import run_agent

# --- Constants ---
MAX_ANALYSES = 3

# --- Example data ---
EXAMPLE_CV = """
John Smith | Data Analyst | john@email.com

EXPERIENCE
Senior Data Analyst - ABC Corp (2021-Present, 3 years)
- Built SQL queries to analyse customer data across 50+ product categories
- Created Power BI dashboards tracking KPIs for operations teams
- Used Python and pandas for data cleaning and ETL pipeline maintenance

Junior Analyst - XYZ Ltd (2019-2021)
- Supported reporting using Excel and SQL
- Automated monthly reporting processes using Python scripts

SKILLS
SQL, Python, pandas, Power BI, Excel, data visualisation

EDUCATION
BSc Statistics - University of Manchester (2019)
"""

EXAMPLE_JD = """
Data Analyst - Logistics Company

We are looking for a Data Analyst with 3+ years of experience.

Required:
- Strong SQL skills
- Python for data analysis (pandas, numpy)
- Power BI or Tableau for dashboards
- Experience with logistics or supply chain data

Nice to have:
- dbt experience
- Azure or AWS
- Freight industry knowledge
"""

EXAMPLE_RESULT = """## Missing Keywords
* numpy
* logistics experience

## Improvement Suggestions
1. Add 'numpy' to your Skills section alongside pandas to match the job description's technical requirements
2. Include 'logistics' or 'supply chain' context in your Experience section
3. Consider adding 'Tableau' to your Skills section as an alternative to Power BI
4. Quantify your Power BI dashboard work — mention number of dashboards, users, or business impact

## Cover Letter
Three years building SQL pipelines and Power BI dashboards for ABC Corp's operations team maps directly onto what you're looking for — the main gap is logistics domain experience, which I'd pick up quickly given my background in operational data.

At ABC Corp, I built complex SQL queries analysing customer purchasing patterns across 50+ categories, created Power BI dashboards that became the primary reporting tool for the operations team, and automated ETL workflows using Python and pandas. The data problems in logistics aren't fundamentally different from what I've been solving — the domain terminology is what I'd need to learn.

Happy to talk through how my experience translates to your freight data needs. You can reach me at john@email.com.
"""

EXAMPLE_COVER_LETTER = """Three years building SQL pipelines and Power BI dashboards for ABC Corp's operations team maps directly onto what you're looking for — the main gap is logistics domain experience, which I'd pick up quickly given my background in operational data.

At ABC Corp, I built complex SQL queries analysing customer purchasing patterns across 50+ categories, created Power BI dashboards that became the primary reporting tool for the operations team, and automated ETL workflows using Python and pandas. The data problems in logistics aren't fundamentally different from what I've been solving — the domain terminology is what I'd need to learn.

Happy to talk through how my experience translates to your freight data needs. You can reach me at john@email.com."""


def score_display(score: int) -> str:
    if score >= 80:
        color = "#52b788"
        label = "Strong Match"
        explanation = "Your CV aligns well with this role. Focus on the missing keywords to maximise your chances."
    elif score >= 60:
        color = "#f4a261"
        label = "Moderate Match"
        explanation = "You meet the core requirements but have some gaps. Address the missing keywords before applying."
    else:
        color = "#e63946"
        label = "Weak Match"
        explanation = "Significant gaps exist between your CV and this role. Consider whether this is the right role to apply for now."

    return f"""
<div style='background: #1a1a2e; padding: 20px; border-radius: 10px;
            text-align: center; margin: 10px 0; border: 2px solid {color}'>
    <div style='font-size: 56px; font-weight: bold; color: {color}'>{score}</div>
    <div style='font-size: 20px; color: #aaa'>/100 — {label}</div>
    <div style='font-size: 14px; color: #888; margin-top: 8px'>{explanation}</div>
</div>
"""


def cover_letter_copy_box(text: str) -> str:
    return f"""
<div style='background: #262730; padding: 15px; border-radius: 5px;
            border: 1px solid #555; font-family: sans-serif;
            font-size: 14px; line-height: 1.6; white-space: pre-wrap;
            word-wrap: break-word; cursor: text; user-select: all;'>
{text}
</div>
"""


def make_docx(text: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading('Cover Letter', 0)
    for para in text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def extract_cv_essentials(cv_text: str, max_chars: int = 3000) -> str:
    """Extract most relevant CV sections for ATS analysis."""
    if len(cv_text) <= max_chars:
        return cv_text

    lines = cv_text.split('\n')
    section_keywords = {
        'skills': ['skill', 'technical', 'competenc', 'expertise', 'proficien'],
        'experience': ['experience', 'employment', 'work history', 'career', 'position'],
        'education': ['education', 'qualification', 'degree', 'academic'],
        'summary': ['summary', 'profile', 'objective', 'about'],
    }

    sections = {}
    current_section = 'other'
    current_lines = []

    for line in lines:
        line_lower = line.lower().strip()
        matched = False
        for section, keywords in section_keywords.items():
            if any(kw in line_lower for kw in keywords) and len(line.strip()) < 50:
                if current_lines:
                    sections[current_section] = sections.get(current_section, []) + current_lines
                current_section = section
                current_lines = [line]
                matched = True
                break
        if not matched:
            current_lines.append(line)

    if current_lines:
        sections[current_section] = sections.get(current_section, []) + current_lines

    priority_order = ['summary', 'skills', 'experience', 'education', 'other']
    result_parts = []
    total_chars = 0

    for section in priority_order:
        if section in sections:
            section_text = '\n'.join(sections[section])
            if total_chars + len(section_text) <= max_chars:
                result_parts.append(section_text)
                total_chars += len(section_text)
            else:
                remaining = max_chars - total_chars
                if remaining > 200:
                    result_parts.append(section_text[:remaining])
                break

    return '\n'.join(result_parts)


def render_result(result: str):
    """Render agent result with score display and cover letter actions."""
    score_match = re.search(r'(\d+)/100', result)
    if score_match:
        score = int(score_match.group(1))
        st.markdown(score_display(score), unsafe_allow_html=True)

    st.markdown(result)

    cover_match = re.search(r'## Cover Letter\n(.*?)(?=##|$)', result, re.DOTALL)
    if cover_match:
        cover_text = cover_match.group(1).strip()
        st.download_button(
            "📥 Download Cover Letter as Word (.docx)",
            data=make_docx(cover_text),
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_{hash(cover_text)}"
        )
        st.caption("Or select all and copy:")
        st.markdown(cover_letter_copy_box(cover_text), unsafe_allow_html=True)


# --- Session state init ---
if 'analysis_count' not in st.session_state:
    st.session_state.analysis_count = 0
if 'jd_text' not in st.session_state:
    st.session_state.jd_text = ''
if 'history' not in st.session_state:
    st.session_state.history = []
if 'show_example' not in st.session_state:
    st.session_state.show_example = False
if 'show_feedback_form' not in st.session_state:
    st.session_state.show_feedback_form = False
if 'current_result' not in st.session_state:
    st.session_state.current_result = None

# --- Page config ---
st.set_page_config(layout="wide", page_title="ATS Gap Analyser")

# --- Sidebar ---
with st.sidebar:
    st.header("How to use")
    st.markdown("""
1. Upload your CV as a PDF or paste the text
2. Paste the job description or fetch from URL
3. Click **Analyse My CV**
4. Get your match score, gaps, and a tailored cover letter
    """)
    st.divider()
    language = st.selectbox(
        "Cover letter language",
        ["English", "Dutch (Nederlands)", "French (Français)", "German (Deutsch)",
         "Spanish (Español)", "Italian (Italiano)", "Portuguese (Português)"],
        index=0,
        help="Analysis stays in English — only the cover letter changes"
    )
    st.divider()
    st.warning(f"⚠️ Demo limit: {MAX_ANALYSES} analyses per session")
    st.divider()
    st.caption("🔒 Your CV and job description are never stored or logged.")
    st.caption("Built with Groq + llama-3.3-70b-versatile")

# --- Header ---
st.title("ATS Gap Analyser")
st.markdown("Paste your CV and a job description to see how well you match and get specific improvement suggestions.")
st.divider()

# --- Inputs ---
col1, col2 = st.columns(2)

with col1:
    st.subheader("Your CV")
    uploaded_file = st.file_uploader("Upload CV as PDF", type="pdf")
  
    if uploaded_file:
        with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
            cv_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        st.success(f"PDF uploaded — {len(cv_text)} characters extracted")
        st.text_area("Extracted CV text", value=cv_text, height=200,
                     disabled=True, label_visibility="visible")
    else:
        cv_text = st.text_area(
            "CV",
            placeholder="Or paste your CV text here...",
            height=250,
            label_visibility="collapsed"
        )

with col2:
    st.subheader("Job Description")

    jd_url = st.text_input(
        "Paste job posting URL (optional)",
        placeholder="https://company.com/careers/job-123"
    )
    # invisible spacer to match URL input height in JD column
    st.markdown("<div style='height: 29px'></div>", unsafe_allow_html=True)

    if jd_url:
        if st.button("📥 Fetch from URL"):
            try:
                import requests
                from bs4 import BeautifulSoup
                with st.spinner("Fetching job description..."):
                    response = requests.get(jd_url, timeout=10, headers={
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                    })
                    soup = BeautifulSoup(response.text, 'html.parser')
                    for tag in soup(['script', 'style', 'nav', 'footer', 'header',
                                     'iframe', 'img', 'button', 'form']):
                        tag.decompose()
                    main_content = (
                        soup.find('main') or
                        soup.find('article') or
                        soup.find(attrs={'class': lambda x: x and any(
                            word in ' '.join(x).lower()
                            for word in ['job', 'posting', 'description', 'content', 'detail']
                        ) if isinstance(x, list) else False}) or
                        soup.find('body')
                    )
                    fetched_text = main_content.get_text(separator='\n', strip=True)[:5000]
                st.session_state.jd_text = fetched_text
                st.success(f"Fetched {len(fetched_text)} characters")
            except Exception as e:
                st.error("Could not fetch this URL — some career portals (Oracle, Workday, SAP) block automated fetching. Please paste the job description text directly.")

    jd_text = st.text_area(
        "JD",
        value=st.session_state.get('jd_text', ''),
        placeholder="Or paste the job description here...",
        height=250,  # always 250, not conditional
        label_visibility="collapsed"
    )

# --- Buttons ---
col_analyse, col_example = st.columns([3, 1])
with col_analyse:
    analyse_button = st.button("🔍 Analyse My CV", type="primary", use_container_width=True)
with col_example:
    if st.button("💡 See example", use_container_width=True):
        st.session_state.show_example = True

# --- Example output ---
if st.session_state.show_example:
    st.divider()
    col_info, col_clear = st.columns([4, 1])
    with col_info:
        st.info("💡 Example output — no API credits used")
    with col_clear:
        if st.button("✕ Clear example"):
            st.session_state.show_example = False
            st.rerun()

    with st.expander("📄 Example CV used"):
        st.text(EXAMPLE_CV)
    with st.expander("📋 Example Job Description used"):
        st.text(EXAMPLE_JD)

    st.markdown(score_display(75), unsafe_allow_html=True)
    st.markdown(EXAMPLE_RESULT)
    st.download_button(
        "📥 Download Cover Letter as Word (.docx)",
        data=make_docx(EXAMPLE_COVER_LETTER),
        file_name="cover_letter.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="example_dl"
    )
    st.caption("Or select all and copy:")
    st.markdown(cover_letter_copy_box(EXAMPLE_COVER_LETTER), unsafe_allow_html=True)

# --- Analyse ---
if analyse_button:
    st.session_state.show_example = False
    st.session_state.current_result = None

    if not cv_text or not jd_text:
        st.warning("Please provide both your CV and the job description.")
    elif len(cv_text.strip()) < 100:
        st.warning("Your CV seems too short. Please paste your full CV text or upload a PDF.")
    elif len(jd_text.strip()) < 50:
        st.warning("The job description seems too short. Please paste the full job description.")
    elif st.session_state.analysis_count >= MAX_ANALYSES:
        st.error(f"You've reached the limit of {MAX_ANALYSES} analyses per session. Please refresh to start a new session.")
    else:
        try:
            cv_input = extract_cv_essentials(cv_text, max_chars=3000)
            jd_input = jd_text[:2000]

            with st.status("Analysing your CV...", expanded=True) as status:
                st.write("Extracting job requirements...")
                result = run_agent(
                        f"Analyse my CV against this job description.\nCV:\n{cv_input}\nJD:\n{jd_input}\nWrite the cover letter in: {language}"
                    )
                status.update(label="Analysis complete!", state="complete", expanded=False)

            st.session_state.analysis_count += 1
            st.session_state.current_result = result

            score_match = re.search(r'(\d+)/100', result)
            st.session_state.history.append({
                'score': score_match.group(0) if score_match else 'N/A',
                'result': result
            })

        except Exception as e:
            if '429' in str(e):
                st.error("Rate limit reached — please try again in a few minutes.")
            else:
                st.error(f"Something went wrong: {str(e)}")

# --- Current result — rendered outside analyse block so download doesn't reset it ---
if st.session_state.current_result:
    st.divider()
    st.subheader(f"Analysis {st.session_state.analysis_count}/{MAX_ANALYSES}")
    render_result(st.session_state.current_result)

# --- History ---
if len(st.session_state.history) > 1:
    st.divider()
    st.subheader("Previous Analyses")
    for i, h in enumerate(st.session_state.history[:-1], 1):
        with st.expander(f"Analysis {i} — Score: {h['score']}"):
            render_result(h['result'])

# --- Feedback ---
if st.session_state.analysis_count > 0:
    st.divider()
    st.caption("Was this analysis helpful?")
    col_up, col_down = st.columns(2)
    with col_up:
        if st.button("👍 Yes, helpful", use_container_width=True):
            logfire.info("user feedback", rating="good",
                        session_count=st.session_state.analysis_count)
            st.success("Thanks!")
    with col_down:
        if st.button("👎 Not helpful", use_container_width=True):
            st.session_state.show_feedback_form = True

    if st.session_state.show_feedback_form:
        feedback_reason = st.selectbox(
            "What was wrong?",
            ['', 'Score seems off', 'Missing keywords wrong',
             'Suggestions not relevant', 'Cover letter too generic',
             'Cover letter had wrong info', 'Other']
        )
        feedback_text = st.text_input("Any other comments? (optional)")
        if st.button("Submit feedback"):
            logfire.info("user feedback",
                        rating="bad",
                        reason=feedback_reason,
                        comment=feedback_text[:200] if feedback_text else "",
                        session_count=st.session_state.analysis_count)
            st.session_state.show_feedback_form = False
            st.info("Thanks — we'll use this to improve.")